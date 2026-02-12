import os
from docx import Document
from markdown_it import MarkdownIt

class PureConverter:
    def __init__(self, template_path: str | None = None):
        """
        初始化转换器
        :param template_path: Word 模板路径 (.docx)
        """
        self.template_path = template_path
        # 初始化 markdown-it，启用 breaks=True 以支持软回车硬换行
        self.md = MarkdownIt('commonmark', {'breaks': True})

    def set_template_path(self, path: str) -> None:
        """Set the Word template path

        Args:
            path: Path to template file (.docx)
        """
        self.template_path = path

    def convert_text(self, md_text: str, settings: dict = {}) -> str:
        """
        预览逻辑 (Convert to String)
        """
        if not md_text:
            return ""
        if settings is None:
            settings = {}

        tokens = self.md.parse(md_text)
        output_lines = []
        list_context = []
        list_counters = []
        
        ignore_bullets = settings.get("ignore_bullets", True)
        ordered_style = settings.get("ordered_list_style", "text")

        idx = 0
        while idx < len(tokens):
            token = tokens[idx]
            
            if token.type == 'bullet_list_open':
                list_context.append('bullet')
            elif token.type == 'ordered_list_open':
                list_context.append('ordered')
                start = token.attrs.get('start', 1) if token.attrs else 1
                list_counters.append(start)
            elif token.type in ['bullet_list_close', 'ordered_list_close']:
                if list_context: list_context.pop()
                if token.type == 'ordered_list_close' and list_counters: list_counters.pop()
            
            elif token.type == 'heading_open':
                level = int(token.tag[1])
                prefix = '#' * level + ' '
                if idx + 1 < len(tokens) and tokens[idx+1].type == 'inline':
                    output_lines.append(f"{prefix}{tokens[idx+1].content}")
                    idx += 1 
                output_lines.append("") 

            elif token.type == 'inline':
                content = token.content
                line = content
                
                if list_context:
                    current_type = list_context[-1]
                    if current_type == 'bullet':
                        if not ignore_bullets:
                            line = f"• {content}"
                    elif current_type == 'ordered':
                        count = list_counters[-1]
                        if ordered_style in ['text', 'list']:
                            line = f"{count}. {content}"
                        list_counters[-1] += 1

                output_lines.append(line)
                
            elif token.type == 'paragraph_close':
                if not list_context:
                    output_lines.append("") 
                
            idx += 1

        return "\n".join(output_lines).strip()

    def convert_to_word(self, md_text: str, output_path: str, settings: dict) -> None:
        """
        导出 Word 文档
        """
        if self.template_path and os.path.exists(self.template_path):
            doc = Document(self.template_path)
            # 清空模板内容，仅保留样式
            for p in doc.paragraphs:
                p._element.getparent().remove(p._element)
            for t in doc.tables:
                t._element.getparent().remove(t._element)
        else:
            doc = Document()

        tokens = self.md.parse(md_text)
        self._render_tokens(doc, tokens, settings)
        doc.save(output_path)

    def _render_tokens(self, doc, tokens, settings):
        """核心渲染逻辑"""
        ignore_bullets = settings.get("ignore_bullets", False)
        ordered_style = settings.get("ordered_list_style", "text")

        idx = 0
        list_stack = []
        ordered_counters = []

        while idx < len(tokens):
            token = tokens[idx]

            # === 1. 列表状态维护 ===
            if token.type == 'bullet_list_open':
                list_stack.append('bullet')
            elif token.type == 'ordered_list_open':
                list_stack.append('ordered')
                start = token.attrs.get('start', 1) if token.attrs else 1
                ordered_counters.append(start)
            elif token.type in ['bullet_list_close', 'ordered_list_close']:
                if list_stack: list_stack.pop()
                if token.type == 'ordered_list_close' and ordered_counters: ordered_counters.pop()

            # === 2. 标题处理 ===
            elif token.type == 'heading_open':
                level = int(token.tag[1])
                if idx + 1 < len(tokens) and tokens[idx+1].type == 'inline':
                    # 标题通常不分行，直接由样式控制
                    p = doc.add_heading('', level=level)
                    self._fill_rich_text(doc, p, tokens[idx+1])
                    idx += 1

            # === 3. 正文/列表项处理 ===
            elif token.type == 'inline':
                parent_type = tokens[idx-1].type
                if parent_type == 'paragraph_open':
                    # 准备前缀和样式
                    p_style = None
                    prefix = ""
                    use_manual_number = False # 标记是否需要手动添加数字

                    if list_stack:
                        curr = list_stack[-1]
                        if curr == 'bullet' and not ignore_bullets:
                            prefix = "• "
                        elif curr == 'ordered':
                            if ordered_style == 'list':
                                # 尝试使用 Word 原生样式
                                p_style = 'List Number'
                            elif ordered_style == 'text':
                                use_manual_number = True
                            # 'none' 什么都不做

                    # === 创建第一段 ===
                    try:
                        # 尝试应用样式 (解决问题2)
                        p = doc.add_paragraph(style=p_style)
                    except KeyError:
                        # 如果模板里没有 'List Number'，回退到普通样式
                        p = doc.add_paragraph()
                        # 并且强制开启手动数字模式
                        if ordered_style == 'list':
                            use_manual_number = True

                    # 统一处理有序列表计数器
                    if list_stack and list_stack[-1] == 'ordered':
                        if use_manual_number:
                            prefix = f"{ordered_counters[-1]}. "
                        ordered_counters[-1] += 1 # 计数增加

                    # 写入前缀
                    if prefix:
                        p.add_run(prefix)

                    # === 核心：渲染富文本并处理换行 (解决问题1) ===
                    # 传入 doc，允许函数内部创建新段落
                    self._fill_rich_text(doc, p, token, style=None) 
                    # 注意：style=None 表示换行后的段落使用默认样式(Normal)
                    # 这样避免换行后的第二行也带上列表编号
            
            idx += 1

    def _fill_rich_text(self, doc, paragraph, inline_token, style=None):
        """
        终极流式渲染 (修复版)：
        同时支持 softbreak 和 hardbreak，确保 breaks=True 时换行生效。
        """
        if not inline_token.children:
            # 处理纯文本 Token (无 children 结构)
            self._add_text_with_breaks(doc, paragraph, inline_token.content, style)
            return

        curr_p = paragraph
        curr_bold = False
        curr_italic = False

        for child in inline_token.children:
            if child.type == 'strong_open':
                curr_bold = True
            elif child.type == 'strong_close':
                curr_bold = False
            elif child.type == 'em_open':
                curr_italic = True
            elif child.type == 'em_close':
                curr_italic = False
            
            # 🟢 [核心修复] 同时捕获软回车(softbreak) 和 硬回车(hardbreak)
            elif child.type == 'softbreak' or child.type == 'hardbreak':
                # 遇到回车 -> 创建真正的 Word 新段落 (硬回车)
                curr_p = doc.add_paragraph(style=style)
            
            elif child.type == 'text' or child.type == 'code_inline':
                # 处理文本内容 (有些文本内部可能还包含 \n)
                parts = child.content.split('\n')
                for i, part in enumerate(parts):
                    if i > 0:
                        # 文本内的 \n 也要分段
                        curr_p = doc.add_paragraph(style=style)
                    
                    if part:
                        run = curr_p.add_run(part)
                        # 关键：新段落继承当前的状态
                        run.bold = curr_bold
                        run.italic = curr_italic
                        if child.type == 'code_inline':
                            run.font.name = 'Courier New'

    def _add_text_with_breaks(self, doc, paragraph, text, style=None):
        """处理无 children 的纯文本换行"""
        lines = text.split('\n')
        curr_p = paragraph
        for i, line in enumerate(lines):
            if i > 0:
                curr_p = doc.add_paragraph(style=style)
            if line:
                curr_p.add_run(line)